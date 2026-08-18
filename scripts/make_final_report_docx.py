"""Dựng báo cáo kết thúc dự án ra file .docx.

    python scripts/make_final_report_docx.py --force
    python scripts/make_final_report_docx.py --out reports/BAO_CAO_CUOI.docx --force

Bảy mục theo đúng yêu cầu: Background · Problem · Method · Result · Conclusion ·
Discussion · Reference. Nội dung, bảng số và hình đều đã điền đầy đủ.

⚠️ **Mọi con số trong file này là hằng số chép tay, không tính lại lúc chạy.** Chúng
được đối chiếu với `src.eval.test_report`, `src.eval.trust` và `src.eval.compare`
chạy trên xác suất đã lưu ở `runs/`, và với `reports/W1..W4_REPORT.md`. Sửa số ở đây
mà không sửa nguồn là tạo drift — đối chiếu lại bằng ba lệnh đó trước khi đổi bất kỳ
ô nào.

Vì sao chép tay thay vì tính lại: bootstrap 2000 lượt trên nhiều cấu hình mất vài
phút và cần `sklearn`; một script sinh tài liệu nên chạy được trong vài giây và
không phụ thuộc kết quả tính lại có trùng tới chữ số cuối hay không. Ba hình thì
ngược lại — chúng sinh từ dữ liệu thật bằng `scripts/make_final_report_figures.py`.

⚠️ Script GHI ĐÈ file đích. Nếu file đã tồn tại thì nó dừng và bắt truyền `--force`.

Kiểu chữ: chỉ dùng font có sẵn trên Windows và phủ đủ dấu tiếng Việt — Cambria cho
phần chữ chạy, Segoe UI cho tiêu đề và bảng. Giống `md2pdf.py` để hai deliverable
nhìn cùng một hệ.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Any

try:
    import docx
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Mm, Pt, RGBColor
except ModuleNotFoundError:  # pragma: no cover - phụ thuộc chỉ dùng ở máy local
    raise SystemExit(
        "Thiếu thư viện python-docx. Cài bằng:\n\n    pip install python-docx\n\n"
        "Nó cố ý không nằm trong requirements.txt — đây là công cụ soạn deliverable "
        "trên máy local, không thuộc môi trường train."
    ) from None


# --- Hằng số trình bày --------------------------------------------------------------

BODY_FONT = "Cambria"
HEAD_FONT = "Segoe UI"
MUTED = RGBColor(0x5C, 0x66, 0x6D)
HEADER_FILL = "EFEFEF"
CALLOUT_FILL = "F5F2E8"
RUO_FOOTER = "Research Use Only — không dùng cho chẩn đoán lâm sàng"

DEFAULT_OUT = Path("reports/FINAL_REPORT.docx")
ASSETS = Path("reports/assets")


# --- Tiện ích XML mà python-docx không bọc sẵn --------------------------------------


def _set_style_font(style: Any, name: str, size_pt: float, *, bold: bool = False) -> None:
    """Đặt font cho một style, kể cả nhánh eastAsia để Word không tự thay chữ."""
    style.font.name = name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def _set_run_font(run: Any, name: str) -> None:
    run.font.name = name
    rfonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def _shade(element: Any, fill: str) -> None:
    """Tô nền một ô bảng hoặc một paragraph (python-docx không có API cho w:shd)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def _repeat_header_row(row: Any) -> None:
    """Đánh dấu hàng tiêu đề để Word lặp lại nó khi bảng tràn trang."""
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _add_field(paragraph: Any, instruction: str) -> None:
    """Chèn field code của Word (PAGE, TOC). Word tự tính lúc mở hoặc khi bấm F9."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, end):
        run._r.append(element)


def _left_bar(paragraph: Any, color: str = "8A6D11") -> None:
    """Vạch dọc bên trái một paragraph — dùng cho khối cảnh báo."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    edge = OxmlElement("w:left")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), "18")
    edge.set(qn("w:space"), "8")
    edge.set(qn("w:color"), color)
    borders.append(edge)
    p_pr.append(borders)


def _emit_runs(paragraph: Any, text: str, size: float, font: str = BODY_FONT) -> None:
    """Ghi text vào paragraph, hiểu `**đậm**` như một đánh dấu inline."""
    for index, chunk in enumerate(text.split("**")):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold = index % 2 == 1
        run.font.size = Pt(size)
        _set_run_font(run, font)


# --- Khối nội dung ------------------------------------------------------------------


class Builder:
    """Gom trạng thái đánh số bảng/hình để các hàm dựng không phải truyền tay."""

    def __init__(self, document: Any) -> None:
        self.doc = document
        self.table_no = 0
        self.figure_no = 0
        self.missing_images: list[str] = []

    def heading(self, number: str, text: str, level: int) -> None:
        """Tiêu đề đánh số THỦ CÔNG — numbering tự động của Word hay trôi khi sửa mục."""
        paragraph = self.doc.add_heading(level=level)
        run = paragraph.add_run(f"{number} {text}" if number else text)
        run.font.color.rgb = RGBColor(0, 0, 0)
        _set_run_font(run, HEAD_FONT)

    def para(self, text: str) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _emit_runs(paragraph, text, 11)

    def bullets(self, items: list[str]) -> None:
        for item in items:
            paragraph = self.doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
            _emit_runs(paragraph, item, 11)

    def callout(self, text: str) -> None:
        """Khối cảnh báo: nền nhạt, vạch dọc bên trái. Dùng cho ràng buộc phải đọc."""
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.3)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(8)
        _left_bar(paragraph)
        _shade(paragraph._p.get_or_add_pPr(), CALLOUT_FILL)
        _emit_runs(paragraph, text, 10)

    def caption(self, text: str) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(10)
        _emit_runs(paragraph, text, 8.5, HEAD_FONT)
        for run in paragraph.runs:
            run.italic = True
            run.font.color.rgb = MUTED

    def table(self, caption: str, headers: list[str], rows: list[list[str]]) -> None:
        """Bảng đã điền số. Cột đầu canh trái, các cột còn lại canh giữa."""
        self.table_no += 1
        head = self.doc.add_paragraph()
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        head.paragraph_format.space_before = Pt(8)
        head.paragraph_format.space_after = Pt(3)
        _emit_runs(head, f"**Bảng {self.table_no}.** {caption}", 8.5, HEAD_FONT)

        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_row = table.rows[0]
        _repeat_header_row(header_row)
        for cell, text in zip(header_row.cells, headers, strict=True):
            _shade(cell._tc.get_or_add_tcPr(), HEADER_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _emit_runs(paragraph, f"**{text}**", 8.5, HEAD_FONT)

        for row, values in zip(table.rows[1:], rows, strict=True):
            for index, (cell, value) in enumerate(zip(row.cells, values, strict=True)):
                paragraph = cell.paragraphs[0]
                if index:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _emit_runs(paragraph, value, 8.5, HEAD_FONT)

        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def figure(self, caption: str, image: Path, width_cm: float = 12.0) -> None:
        """Chèn hình. Thiếu tệp thì để khung trống chứ không nổ — báo cáo vẫn dựng được."""
        self.figure_no += 1
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(3)
        if image.exists():
            paragraph.add_run().add_picture(str(image), width=Cm(width_cm))
        else:
            self.missing_images.append(str(image))
            run = paragraph.add_run(f"[Thiếu tệp hình: {image}]")
            run.italic = True
            run.font.color.rgb = MUTED
            run.font.size = Pt(10)
        self.caption(f"**Hình {self.figure_no}.** {caption}")

    def page_break(self) -> None:
        self.doc.add_page_break()


# ====================================================================================
# NỘI DUNG
# ====================================================================================
#
# Khai báo dạng DỮ LIỆU rồi một vòng lặp dựng — sửa báo cáo là sửa danh sách này,
# không đụng logic ở trên. Mỗi phần tử là một tuple ("loại", ...):
#     ("h", số, tiêu đề, cấp) · ("p", văn xuôi) · ("ul", [gạch đầu dòng])
#     ("callout", cảnh báo) · ("table", caption, [cột], [[hàng]])
#     ("fig", caption, tên tệp ảnh, bề rộng cm) · ("pagebreak",)
# Trong văn xuôi, `**...**` là in đậm.

CONTENT: list[tuple[Any, ...]] = [
    # ============================== 1. BACKGROUND ==============================
    ("h", "1.", "Background", 1),
    ("h", "1.1.", "Bối cảnh lâm sàng", 2),
    (
        "p",
        "Tổn thương khu trú ở gan bao gồm nhiều loại có tiên lượng và hướng xử trí khác "
        "nhau rất xa. Nang gan và u máu phần lớn là phát hiện tình cờ và thường không cần "
        "can thiệp; áp-xe cần điều trị nhiễm khuẩn; tăng sản thể nốt khu trú (FNH) là tổn "
        "thương lành cần theo dõi; trong khi ung thư biểu mô tế bào gan (HCC), ung thư "
        "đường mật trong gan (ICC) và tổn thương di căn đều là bệnh ác tính với phác đồ "
        "khác nhau. **Gọi đúng tên tổn thương vì vậy quyết định trực tiếp hướng xử trí "
        "tiếp theo**, chứ không chỉ là một nhãn mô tả.",
    ),
    (
        "p",
        "Cộng hưởng từ có thuốc tương phản là phương tiện hình ảnh giàu thông tin nhất cho "
        "việc phân biệt này, vì nó không chỉ cho hình thái tổn thương mà còn cho **động "
        "học ngấm thuốc theo thời gian**. Chính chuỗi thay đổi tín hiệu qua các thì trước "
        "thuốc, động mạch, tĩnh mạch cửa và thì muộn mới là dấu hiệu phân biệt: HCC điển "
        "hình ngấm mạnh thì động mạch rồi thải thuốc, ICC ngấm tiến triển, tổn thương di "
        "căn thường có viền ngấm. Đọc đúng vì thế đòi hỏi so sánh nhiều chuỗi ảnh với "
        "nhau, chứ không phải đọc từng chuỗi riêng lẻ.",
    ),
    (
        "p",
        "Đây cũng chính là chỗ bài toán trở nên khó cho cả người đọc lẫn máy: tín hiệu "
        "phân biệt nằm ở **quan hệ giữa các thì**, mà các thì lại được chụp ở những lần "
        "nín thở khác nhau nên không khớp nhau về mặt không gian. Phần Method sẽ cho thấy "
        "đây không phải một nhận xét lý thuyết — nó là nguyên nhân đo được của phần lớn "
        "khoảng cách hiệu năng trong nửa đầu dự án.",
    ),
    ("h", "1.2.", "Bối cảnh kỹ thuật và công trình liên quan", 2),
    (
        "p",
        "Bài toán này có một mốc đối chiếu công khai rất tốt: **challenge LLD-MMRI tại "
        "MICCAI 2023**, với tập test 104 ca được ban tổ chức giữ kín và một bảng xếp hạng "
        "công bố. Baseline chính thức của ban tổ chức là UniFormer-S 3D huấn luyện từ đầu, "
        "đạt macro-F1 0,6083. Ba đội dẫn đầu đạt 0,8322 · 0,8078 · 0,7860. Sau challenge, "
        "một loạt phương pháp được công bố trên đúng tập test đó, trong đó SDR-Former "
        "0,791, STM-Former 0,793, RadioFormer 0,806 và CGHNet 0,818.",
    ),
    (
        "p",
        "Điểm chung của toàn bộ nhóm công trình này là **chúng chỉ báo cáo năng lực phân "
        "loại** — accuracy, macro-F1, Cohen's κ, đôi khi AUC. Không công trình nào trong "
        "số đó báo cáo xác suất của mô hình có đáng tin hay không, và không công trình nào "
        "đo xem mô hình có biết lúc nào nó đang không chắc hay không. Với một công cụ dự "
        "định đặt cạnh bác sĩ, đó là khoảng trống đáng kể: một mô hình sai 20% số ca nhưng "
        "biết chỉ ra 20% ca nào đáng ngờ thì dùng được, còn một mô hình sai 20% số ca một "
        "cách tự tin thì không.",
    ),
    (
        "p",
        "**Dự án này nhắm vào khoảng trống đó**, chứ không nhắm vào việc leo bảng xếp "
        "hạng. Lý do thứ hai cũng thuần tuý thống kê: với n = 104, khoảng tin cậy bootstrap "
        "rộng khoảng ±0,08 quanh mức 0,83, nên hai phương pháp cách nhau vài điểm phần trăm "
        "không thể phân biệt được trên tập test này. Đua điểm ở cỡ mẫu đó là đua với nhiễu.",
    ),
    ("h", "1.3.", "Lịch sử phạm vi của dự án", 2),
    (
        "p",
        "Dự án khởi đầu ở một bài toán khác: phân loại nhị phân trên CT, dùng LiTS để suy "
        "nhãn ở mức lát cắt (lát có tổn thương gan so với lát gan bình thường) rồi gộp lên "
        "mức bệnh nhân. Pha đó đã chạy được và tạo ra nền tảng kỹ thuật thật: audit 131 "
        "volume, chia tập ở mức bệnh nhân có kiểm tra rò rỉ, bootstrap CI, sàng lọc bốn "
        "backbone, và một lần đánh giá ngoài trên 3D-IRCADb-01 (slice-AUROC 0,807 "
        "[0,678; 0,902]).",
    ),
    (
        "p",
        "Nhưng nhãn của LiTS được suy ra từ mặt nạ phân vùng, nên bài toán nhị phân đó chỉ "
        "trả lời một câu hỏi thay thế — *có tổn thương hay không* — chứ không trả lời câu "
        "hỏi lâm sàng thật là *đây là tổn thương loại gì*. Sau phản hồi của người hướng dẫn "
        "và một vòng rà soát độ khớp giữa dữ liệu, nhãn và câu hỏi nghiên cứu, dự án "
        "chuyển hẳn sang **phân loại đa lớp trên MRI 3D đa thì với LLD-MMRI**, và lấy độ "
        "tin cậy của xác suất làm đóng góp chính. Đây là quyết định tăng độ khớp giữa dữ "
        "liệu và câu hỏi, không phải phủ nhận giá trị của pha CT.",
    ),
    ("pagebreak",),
    # =============================== 2. PROBLEM ================================
    ("h", "2.", "Problem", 1),
    ("h", "2.1.", "Phát biểu bài toán", 2),
    (
        "p",
        "Cho trước một bộ ảnh MRI gan đa thì của một bệnh nhân và **vị trí tổn thương đã "
        "được khoanh sẵn**, hãy xác định tổn thương đó thuộc một trong bảy loại: HCC, ICC, "
        "di căn, nang, u máu, FNH, áp-xe. Ba loại đầu là ác tính, bốn loại sau là lành tính.",
    ),
    (
        "p",
        "Hai điều bài toán này **không** phải, và cần nói rõ để tránh hiểu nhầm phạm vi. "
        "Nó không phải bài toán phát hiện: vị trí tổn thương được cho trước bằng khung bao "
        "có sẵn trong annotation, mô hình không phải đi tìm. Nó cũng không phải phân vùng: "
        "đầu ra là một nhãn cho một tổn thương, không phải một mặt nạ theo voxel.",
    ),
    ("h", "2.2.", "Đầu vào và đầu ra", 2),
    (
        "p",
        "**Đầu vào:** tám thì MRI của cùng một bệnh nhân — C-pre (trước thuốc), C+A (động "
        "mạch), C+V (tĩnh mạch cửa), C+Delay (thì muộn), T2WI, DWI, In Phase và Out Phase "
        "— được cắt quanh tổn thương và xếp thành tám kênh của một khối 3D duy nhất, kích "
        "thước 8 × 112 × 112 × 14.",
    ),
    (
        "p",
        "**Đầu ra gồm ba thứ, và hai thứ sau mới là đóng góp của đề tài:** (1) bảy xác "
        "suất, một cho mỗi lớp; (2) một **mức bất định** cho biết mô hình đang chắc đến đâu "
        "ở chính ca này; (3) một **cờ từ chối** — khi mức bất định vượt ngưỡng, hệ thống "
        "không đưa ra câu trả lời mà chuyển ca đó cho bác sĩ đọc. Mức bất định và cờ từ "
        "chối không phải phụ kiện trang trí cho xác suất; chúng là thứ quyết định hệ thống "
        "có dùng được cạnh một người đọc hay không.",
    ),
    ("h", "2.3.", "Câu hỏi nghiên cứu và đóng góp", 2),
    (
        "p",
        "Câu hỏi chính: **một mô hình phân loại tổn thương gan trên MRI đa thì, huấn luyện "
        "trên chưa tới 400 bệnh nhân, có thể đưa ra xác suất đáng tin và biết từ chối ca nó "
        "không chắc hay không?** Cụ thể hoá thành hai câu đo được:",
    ),
    (
        "ul",
        [
            "**Calibration.** Khi mô hình nói 80% thì nó có đúng khoảng 80% số lần hay "
            "không? Đo bằng ECE, MCE, Brier và NLL, có so với bản đã hiệu chỉnh nhiệt độ.",
            "**Selective prediction.** Nếu cho phép mô hình từ chối một phần số ca, phần "
            "nó giữ lại có chính xác lên thật hay không, và lên bao nhiêu? Đo bằng đường "
            "risk–coverage, AURC, và macro-F1 tại các mức coverage, mỗi mức kèm kiểm định "
            "ghép cặp so với trường hợp không từ chối ca nào.",
        ],
    ),
    (
        "p",
        "Hai câu hỏi phụ trợ đi kèm: **nguồn khởi tạo trọng số** có ảnh hưởng bao nhiêu ở "
        "cỡ dữ liệu này, và **cách chuẩn bị dữ liệu đa thì** đóng góp bao nhiêu so với việc "
        "đổi kiến trúc. Cả hai đều được trả lời bằng số trong phần Result.",
    ),
    ("h", "2.4.", "Phạm vi và ràng buộc", 2),
    (
        "ul",
        [
            "**Research Use Only.** Kết quả chưa qua bất kỳ kiểm định lâm sàng nào và "
            "không được dùng để chẩn đoán hay thay thế bác sĩ.",
            "**Nguồn lực:** một người thực hiện, sáu tuần.",
            "**Compute:** Kaggle Notebook, mỗi phiên tối đa 12 giờ và quota 30 giờ mỗi "
            "tuần. Ràng buộc này định hình thiết kế code ngay từ đầu, không phải sửa sau.",
            "**Ba loại sản phẩm bàn giao:** mã tiền xử lý / huấn luyện / đánh giá; một web "
            "app demo tự viết full-stack; và slide cùng báo cáo.",
        ],
    ),
    ("pagebreak",),
    # ================================ 3. METHOD =================================
    ("h", "3.", "Method", 1),
    ("h", "3.1.", "Dữ liệu", 2),
    ("h", "3.1.1.", "Dataset LLD-MMRI", 3),
    (
        "p",
        "Dataset gồm **498 bệnh nhân**, mỗi bệnh nhân đúng một tổn thương, mỗi tổn thương "
        "có tám thì MRI — tổng cộng 3.984 khối ảnh 3D. Mức mất cân bằng lớp là vừa phải: "
        "tỉ lệ giữa lớp đông nhất (HCC) và lớp hiếm nhất (FNH) là khoảng 3,4 : 1. Đây "
        "không phải phân bố đuôi dài, nên các kỹ thuật dành riêng cho long-tail không phải "
        "công cụ đúng ở đây.",
    ),
    (
        "table",
        "Phân bố bảy lớp trên toàn bộ dataset (n = 498).",
        ["Lớp", "Tên đầy đủ", "Nhóm", "n"],
        [
            ["u máu", "Hepatic hemangioma", "lành", "79"],
            ["ICC", "Intrahepatic cholangiocarcinoma", "ác", "58"],
            ["áp-xe", "Hepatic abscess", "lành", "54"],
            ["di căn", "Hepatic metastasis", "ác", "51"],
            ["nang", "Hepatic cyst", "lành", "53"],
            ["FNH", "Focal nodular hyperplasia", "lành", "46"],
            ["HCC", "Hepatocellular carcinoma", "ác", "157"],
            ["**Tổng**", "", "3 ác / 4 lành", "**498**"],
        ],
    ),
    ("h", "3.1.2.", "Tám thì MRI và đặc điểm hình học", 3),
    (
        "p",
        "Một gate kiểm tra hình học chạy trên toàn bộ dữ liệu thật đã làm lộ ra hai điều "
        "không lường trước, và cả hai về sau đều ảnh hưởng trực tiếp tới thiết kế:",
    ),
    (
        "ul",
        [
            "**Tám thì của cùng một bệnh nhân không nằm trên cùng một lưới voxel.** Có ca "
            "thì động là 512 × 512 × 88 với lát dày 2,6 mm trong khi T2WI là 512 × 512 × 24 "
            "với lát dày 9 mm. Khác mặt phẳng chụp, khác ma trận, và dữ liệu đến từ nhiều "
            "máy 1.5T lẫn 3T.",
            "**Nhóm In/Out Phase không cố định.** Phần lớn ca chúng đi cùng nhóm thì động, "
            "nhưng một số ca lại đi cùng nhóm T2WI. Hệ quả thiết kế: không thể cố định "
            "nhóm thì khi làm cơ chế hợp nhất.",
        ],
    ),
    (
        "p",
        "Quan trọng nhất là con số thứ ba. Tâm tổn thương của cùng một bệnh nhân, đo trong "
        "toạ độ thế giới, **tán ra 23,3 mm theo trục đầu–chân** giữa tám thì. Đây không "
        "phải lỗi đo mà là biên độ chuyển động hô hấp của gan, vì tám thì được chụp ở các "
        "lần nín thở khác nhau. Trên một cửa sổ cắt chỉ khoảng 43,6 mm thì độ lệch đó "
        "chiếm hơn nửa chiều sâu.",
    ),
    ("h", "3.1.3.", "Nhãn, mặt nạ và gold standard", 3),
    (
        "p",
        "Nhãn bảy lớp lấy từ **kết quả giải phẫu bệnh**, không phải từ đọc ảnh — đây là "
        "gold standard thật và là một điểm mạnh của dataset này. Kèm theo là khung bao 2D "
        "theo từng lát cho vị trí tổn thương.",
    ),
    (
        "p",
        "Bộ mặt nạ phân vùng đi kèm là nhãn chính thức của LLD-MMRI, bổ sung tháng 3/2025, "
        "gán bằng MedSAM2 trong một quy trình có người trong vòng lặp. Dự án dùng nó để "
        "xác định cửa sổ cắt bám tổn thương. Cần giữ một dè dặt có cơ sở: mức can thiệp "
        "của con người trong quy trình đó không được mô tả rõ, nên dùng mặt nạ làm mục "
        "tiêu phụ trợ thì hợp lý, còn báo cáo chất lượng phân vùng như một kết quả thì phải "
        "nêu giới hạn.",
    ),
    ("h", "3.2.", "Chia dữ liệu", 2),
    ("h", "3.2.1.", "Split official 316 / 78 / 104", 3),
    (
        "p",
        "Bản dataset thực nhận **không kèm file chia tập**. Quyết định ban đầu là tự chia "
        "5-fold phân tầng; quyết định đó đã bị đảo ngược sau khi tìm được danh sách 394 ca "
        "train+val mà một đội dự thi công khai, từ đó suy ra tập test = 498 − 394. Phân bố "
        "lớp của bản tái lập được đối chiếu với tài liệu chính thức của challenge và "
        "**khớp đủ 7/7 lớp**, nên split 316 / 78 / 104 đã được khôi phục đúng.",
    ),
    (
        "p",
        "Đây là quyết định quan trọng nhất về mặt phương pháp trong cả dự án. Tự chia thì "
        "mọi con số về sau chỉ so được với chính mình; dùng đúng split chính thức thì "
        "**so trực tiếp được với bảng xếp hạng và với văn liệu**. Các file split đã được "
        "commit và khoá lại; quality gate chặn mọi thay đổi lên thư mục đó.",
    ),
    (
        "table",
        "Phân bố lớp theo tập sau khi tái lập split chính thức. Đã đối chiếu khớp 7/7 lớp "
        "với tài liệu challenge.",
        ["Lớp", "train + val", "test", "Tổng"],
        [
            ["u máu", "63", "16", "79"],
            ["ICC", "46", "12", "58"],
            ["áp-xe", "42", "12", "54"],
            ["di căn", "40", "11", "51"],
            ["nang", "42", "11", "53"],
            ["FNH", "36", "10", "46"],
            ["HCC", "125", "32", "157"],
            ["**Tổng**", "**394**", "**104**", "**498**"],
        ],
    ),
    ("h", "3.2.2.", "Cross-validation 5 fold", 3),
    (
        "p",
        "Việc chọn mô hình dùng 5 fold chính thức trên 394 ca train+val. Năm tập validation "
        "phân hoạch sạch 394 ca — đã kiểm chứng trực tiếp rằng giao của mọi cặp bằng rỗng "
        "và hợp của cả năm đúng bằng 394.",
    ),
    (
        "table",
        "Cỡ tập train và validation của từng fold.",
        ["Fold", "n train", "n val"],
        [
            ["1", "312", "82"],
            ["2", "314", "80"],
            ["3", "316", "78"],
            ["4", "317", "77"],
            ["5", "317", "77"],
            ["**Tổng out-of-fold**", "—", "**394**"],
        ],
    ),
    ("h", "3.2.3.", "Nguyên tắc chống rò rỉ dữ liệu", 3),
    (
        "ul",
        [
            "**Chia tập tuyệt đối ở mức bệnh nhân.** Không bao giờ để hai thì của cùng một "
            "bệnh nhân rơi vào hai tập khác nhau. Có unit test kiểm giao tập bệnh nhân "
            "bằng rỗng, chạy trong quality gate.",
            "**Thống kê chuẩn hoá chỉ tính trong phạm vi từng bệnh nhân**, không gộp xuyên "
            "bệnh nhân.",
            "**Không chọn mô hình, ngưỡng hay nhiệt độ trên tập test.** Nhiệt độ hiệu chỉnh "
            "được fit trên out-of-fold rồi áp mù lên test.",
            "**Nhiệt độ trên out-of-fold fit theo kiểu leave-one-fold-out**: giá trị áp lên "
            "một fold học từ bốn fold còn lại, nên không ca nào được hiệu chỉnh bởi một "
            "nhiệt độ đã nhìn thấy nó.",
            "**Không gộp 5 checkpoint của 5 fold để báo số out-of-fold.** Mỗi ca ở "
            "validation của một fold nằm trong tập train của bốn mô hình còn lại. Phép gộp "
            "đó chỉ hợp lệ trên tập test, nơi không mô hình nào từng nhìn thấy ca nào.",
        ],
    ),
    ("h", "3.3.", "Tiền xử lý", 2),
    ("h", "3.3.1.", "Gate hình học và phán quyết thứ tự trục", 3),
    (
        "p",
        "Rủi ro cụ thể cần loại trừ trước khi làm bất cứ việc gì khác: bản dataset có thể "
        "đã lấy mẫu lại hoặc xoay hướng ảnh trong khi khung bao vẫn ở toạ độ gốc. Nếu vậy "
        "thì mọi patch cắt theo khung bao đều lệch, **và không có gì trong quá trình huấn "
        "luyện báo hiệu điều đó**. Gate đối chiếu spacing trong header ảnh với spacing khai "
        "trong annotation, cộng kiểm chỉ số lát và kiểm khung bao có nằm trong biên. "
        "Kết quả: **đạt 3.984/3.984**.",
    ),
    (
        "p",
        "Một vấn đề nữa: annotation không nói rõ khung bao là (x, y) hay (y, x). Mọi ảnh "
        "đều vuông nên khung bao lọt cả hai cách hiểu, mà đoán sai thì mọi patch lệch 90°. "
        "Cách giải: cùng một tổn thương vật lý thì tâm khung bao của tám thì phải hội tụ "
        "trong toạ độ thế giới, còn cách hiểu sai sẽ làm chúng tán ra. Đo trong mặt phẳng "
        "cho phán quyết rõ ràng — **166/180 phiếu (92%) cho cách hiểu (x, y)**, với độ tán "
        "trung vị 12,4 mm so với 17,8 mm của cách hiểu còn lại.",
    ),
    ("h", "3.3.2.", "Cắt trong không gian mm và căn các thì", 3),
    (
        "p",
        "Vì tám thì khác lưới, khung bao tính theo voxel của thì này vô nghĩa với thì kia; "
        "nhưng cả tám chung một hệ toạ độ bệnh nhân. Cách làm: đổi tâm khung bao sang mm, "
        "dựng một lưới đích chung quanh tâm đó, rồi lấy mẫu cả tám thì lên lưới ấy. Cửa sổ "
        "cắt **bám sát tổn thương** — lấy khối tổn thương theo mặt nạ, nhân hệ số viền 1,6, "
        "chặn trong khoảng 40–200 mm — rồi đưa về kích thước đích.",
    ),
    (
        "p",
        "Bước quyết định là **căn từng thì về tâm tổn thương của chính nó**, dùng toạ độ có "
        "sẵn trong annotation. Spacing và trường nhìn tính một lần từ thì tham chiếu (C+V), "
        "chỉ đổi tâm, nên tám khối giữ nguyên cùng kích thước vật lý và khác nhau đúng một "
        "phép tịnh tiến. Chi phí chỉ là một lần dựng lại cache, không cần thuật toán "
        "registration.",
    ),
    (
        "callout",
        "⚠️ **Đây không phải một phép sửa trung tính, và điều đó phải được nói rõ.** Phép "
        "căn này chỉ khử tịnh tiến; nó không khử xoay và không khử biến dạng. Hệ quả là mô "
        "xung quanh sẽ thôi khớp giữa các thì, chỉ riêng tổn thương khớp. Với một bài toán "
        "phân loại tổn thương thì đó có thể là điều mong muốn, nhưng nó là một thay đổi về "
        "ngữ nghĩa của dữ liệu, không chỉ là một bước tiền xử lý.",
    ),
    (
        "p",
        "Một gate chạy trước khi huấn luyện đã kiểm rằng phép căn thật sự có hiệu lực — nếu "
        "không thì cache mới sẽ giống hệt cache cũ và bốn giờ huấn luyện sẽ lặp lại đúng "
        "kết quả cũ mà đường cong không hé lộ gì. Gate qua: cả 498 ca cắt theo mặt nạ, "
        "không ca nào phải lùi về tâm tham chiếu, độ dịch giữa các thì có **trung vị "
        "19,65 mm** (nhỏ nhất 2,80, lớn nhất 111,0).",
    ),
    ("h", "3.3.3.", "Chuẩn hoá cường độ", 3),
    (
        "p",
        "MRI không có đơn vị chuẩn như HU của CT, nên chuẩn hoá là bắt buộc. Cách làm: cắt "
        "đuôi ở phân vị 0,5 và 99,5 rồi z-score, với thống kê tính trên **toàn bộ volume "
        "gốc của chính bệnh nhân đó** trước khi cắt patch. Vì thống kê không gộp qua nhiều "
        "bệnh nhân nên bước này không vi phạm nguyên tắc chống rò rỉ.",
    ),
    ("h", "3.3.4.", "Các bản cache đã dựng", 3),
    (
        "p",
        "Tiền xử lý chạy một lần rồi lưu thành cache và đẩy lên Kaggle Dataset có đánh "
        "phiên bản; notebook huấn luyện chỉ mount vào. Ba bản cache tồn tại vì ba hình học "
        "đầu vào khác nhau đã được thử. Ở cả ba, spacing suy từ kích thước **trong** (phần "
        "mô hình nhận) chứ không từ lưới cache, nên độ phân giải vật lý giữ nguyên và cắt "
        "giữa cache có lề cho ra đúng khối mà cache không lề tạo ra.",
    ),
    (
        "table",
        "Ba bản cache tiền xử lý.",
        ["Cache", "Kích thước trong", "Lề mỗi phía", "Mảng lưu", "Dung lượng", "Dùng cho"],
        [
            [
                "E4",
                "112 × 112 × 32",
                "không",
                "8 × 112 × 112 × 32",
                "≈ 3,2 GB",
                "E0–E6b, web app",
            ],
            [
                "E12",
                "112 × 112 × 32",
                "12 / 12 / 4",
                "8 × 136 × 136 × 40",
                "≈ 5,9 GB",
                "E12 (cắt ngẫu nhiên)",
            ],
            [
                "CGHNet",
                "112 × 112 × 14",
                "8 / 8 / 1",
                "8 × 128 × 128 × 16",
                "≈ 2,1 GB",
                "CGHNet, **UniFormer**",
            ],
        ],
    ),
    ("h", "3.4.", "Kiến trúc mô hình", 2),
    (
        "p",
        "Bốn kiến trúc đã được huấn luyện thật. Điểm khác biệt quyết định của cấu hình cuối "
        "cùng **không nằm ở kiến trúc mà ở nguồn khởi tạo trọng số**: UniFormer-S 3D được "
        "nạp trọng số học trên Kinetics-400, một tập dữ liệu video. Mô hình vì thế bắt đầu "
        "từ một biểu diễn đã quen với cấu trúc ba chiều, thay vì từ trọng số ngẫu nhiên.",
    ),
    (
        "p",
        "Đây là chỗ khác biệt so với các lần thử trước trong dự án. Trước đó mô hình hoặc "
        "huấn luyện từ đầu, hoặc dùng trọng số học trên ảnh 2D. Pre-training 2D không học "
        "trực tiếp quan hệ theo chiều sâu của khối, nên không tạo được lợi ích rõ ràng. "
        "Đáng chú ý: **baseline chính thức của challenge dùng đúng kiến trúc UniFormer-S "
        "3D nhưng huấn luyện từ đầu và chỉ đạt 0,6083**, nên đây gần như là một phép so "
        "sánh có kiểm soát sẵn có trong văn liệu.",
    ),
    (
        "table",
        "Các kiến trúc đã huấn luyện và đánh giá.",
        ["Kiến trúc", "Khởi tạo", "Hình học vào", "Tham số", "Trạng thái"],
        [
            [
                "DenseNet121-3D",
                "ngẫu nhiên",
                "112 × 112 × 32",
                "11,4 M",
                "mốc đối chứng, đã đánh giá trên test",
            ],
            [
                "Siamese đa thì",
                "ngẫu nhiên",
                "96 × 96 × 48",
                "—",
                "huỷ — lẫn biến gây nhiễu",
            ],
            [
                "CGHNet (tái lập)",
                "ngẫu nhiên",
                "112 × 112 × 14",
                "—",
                "đủ 5 fold, kết quả âm",
            ],
            [
                "**UniFormer-S 3D**",
                "**Kinetics-400**",
                "112 × 112 × 14",
                "≈ 21 M *",
                "**cấu hình chính**",
            ],
        ],
    ),
    (
        "p",
        "* Con số tham số của UniFormer là của bản gốc dành cho video; nó chưa được đo "
        "lại sau khi đổi lớp vào sang 8 kênh và đầu ra sang 7 lớp. Ngoài bốn kiến trúc "
        "trên, repo còn giữ cấu hình của UniFormer-Base, UniFormerV2-B/16 và SDR-Former "
        "nhưng **chưa chạy fold nào** — chúng thuộc phần Hướng phát triển.",
    ),
    ("h", "3.5.", "Cấu hình huấn luyện", 2),
    (
        "p",
        "Mọi siêu tham số đi qua file YAML trong `configs/`, không hardcode trong mã huấn "
        "luyện, và seed cố định ở 1337. Cấu hình chính là bản **tái lập recipe của đội hạng "
        "hai** trong challenge, có chú thích nguồn cho từng khoá. Năm lần huấn luyện của 5 "
        "fold dùng cùng một file cấu hình, khác nhau đúng khoá `fold` — điều này được kiểm "
        "chứng bằng cách so từng khoá của năm bản cấu hình đã lưu, **trước khi** đọc kết quả.",
    ),
    (
        "table",
        "Siêu tham số của hai cấu hình.",
        ["Siêu tham số", "DenseNet121-3D (đối chứng)", "UniFormer-S + Kinetics (chính)"],
        [
            ["Số epoch", "300", "300"],
            ["Optimizer", "AdamW", "AdamW"],
            ["Learning rate", "1e-4", "1e-4"],
            ["Weight decay", "0,05", "0,05"],
            ["Warmup", "5 epoch, lr 1e-6", "5 epoch, lr 1e-6"],
            ["Lịch learning rate", "cosine, min 1e-5", "cosine, min 1e-5"],
            ["Batch size", "2", "4"],
            ["Batch hiệu dụng", "8 (accum 4)", "4 (accum 1)"],
            ["Hàm mất mát", "CrossEntropy", "Focal, γ = 2 (softmax)"],
            ["Trọng số lớp", "không", "số mẫu hiệu dụng, β = 0,9999"],
            ["Label smoothing", "0", "0,1"],
            ["Bộ lấy mẫu", "đều", "căn bậc hai theo tần suất lớp"],
            ["Drop-path", "—", "0,1"],
            ["Dropout ở đầu ra", "0,2", "0,0"],
            ["AMP", "có", "có"],
            ["Early stopping", "không", "không"],
            ["Seed", "1337", "1337"],
        ],
    ),
    (
        "callout",
        "⚠️ **Bốn chỗ cố ý lệch khỏi recipe gốc, đều phải tính đến khi đọc kết quả.** "
        "(1) Focal loss của dự án là bản **softmax**, của họ là bản **sigmoid** — đây là "
        "chỗ lệch đáng kể nhất. (2) Dùng biến thể `small` thay vì `base`, vì chỉ bản small "
        "có đúng tệp trọng số công khai. (3) Phép xoay lấp biên bằng cách nhân bản voxel "
        "biên thay vì lấp 0 — chỗ này dự án làm đúng hơn họ. (4) **Cấu hình chính không "
        "bật phép nội suy hai ca cùng chẩn đoán** mà recipe của họ có dùng. Vì tái lập "
        "không trọn vẹn, khoảng cách 0,039 điểm so với kết quả của họ nằm trong phạm vi "
        "bốn chỗ lệch này cộng nhiễu cỡ mẫu.",
    ),
    ("h", "3.6.", "Tăng cường dữ liệu", 2),
    (
        "p",
        "Ba phép biến đổi kết cấu (edge, emboss, nhóm lọc) **loại trừ lẫn nhau** theo một "
        "cây quyết định, nên tổng lại **60% số mẫu huấn luyện không bị phép nào trong nhóm "
        "này**. Nhẹ hơn nhiều so với việc bật cả ba. Mọi phép đều áp cùng một tham số cho "
        "cả tám thì — điều này quan trọng, xem mục 4.9.",
    ),
    (
        "table",
        "Tăng cường dữ liệu của cấu hình chính.",
        ["Phép biến đổi", "Tham số", "Xác suất", "Ghi chú"],
        [
            ["Lật theo trục x / y / z", "mỗi trục độc lập", "0,5", "—"],
            ["Xoay trong mặt phẳng", "±10°", "1,0", "lấp biên bằng voxel biên"],
            ["Cắt ngẫu nhiên", "128 → 112 và 16 → 14", "1,0", "biên độ ±7,1%"],
            ["Nhiễu cường độ", "—", "0", "**tắt** — đã đo là gây hại"],
            ["Edge", "—", "0,10", "ba nhóm loại trừ nhau"],
            ["Emboss", "—", "0,10", "ba nhóm loại trừ nhau"],
            ["Blur", "σ = 1", "0,08", "trong nhóm lọc"],
            ["Sharpen", "—", "0,08", "trong nhóm lọc"],
            ["Unsharp", "—", "0,04", "trong nhóm lọc"],
            ["**Không phép nào**", "—", "**0,60**", "—"],
        ],
    ),
    ("h", "3.7.", "Danh sách thí nghiệm", 2),
    (
        "p",
        "Chuỗi thí nghiệm được thiết kế để mỗi lần chỉ đổi một cụm biến. Một quy tắc đã "
        "được chốt sau khi dự án bị đánh lừa vài lần: **một phép sàng trên cỡ mẫu nhỏ chỉ "
        "đủ để loại một ý tưởng, không đủ để chọn nó**. Bằng chứng cho quy tắc đó nằm ở "
        "mục 6.3.",
    ),
    (
        "table",
        "Toàn bộ thí nghiệm đã chạy, theo thứ tự thời gian.",
        ["Mã", "Thay đổi so với mốc trước", "Số fold", "Kết luận"],
        [
            ["E0", "tái lập recipe chính thức, cửa sổ mm cố định", "1", "mốc xuất phát"],
            ["E1", "cache cắt bám sát tổn thương", "1", "tăng rõ"],
            ["E2", "Siamese đa thì, trọng số dùng chung", "—", "huỷ, lẫn biến"],
            ["E3", "hình học 112 × 112 × 32", "1", "dừng sớm, không kết luận"],
            ["E4", "căn từng thì theo tổn thương của chính nó", "5", "**cấu hình đối chứng**"],
            ["E5", "focal loss γ = 2", "2", "không có ý nghĩa thống kê"],
            ["E6", "tăng cường hình học mạnh hơn + nhiễu cường độ", "2", "không có ý nghĩa"],
            ["E6b", "như E6 nhưng tắt nhiễu cường độ", "5", "null trên đủ 5 fold"],
            ["E12", "cắt ngẫu nhiên từ cache rộng hơn", "3", "null"],
            ["TTA", "test-time augmentation bằng phép lật", "5", "âm"],
            ["CGHNet", "tái lập kiến trúc lai 2D–3D", "5", "âm"],
            [
                "**UniFormer**",
                "**UniFormer-S 3D + trọng số Kinetics-400**",
                "**5**",
                "**cấu hình chính**",
            ],
        ],
    ),
    ("h", "3.8.", "Phương pháp đánh giá", 2),
    ("h", "3.8.1.", "Metric", 3),
    (
        "ul",
        [
            "**Macro-F1** là metric chốt. F1 tính riêng cho từng lớp rồi lấy trung bình "
            "đều, nên **lớp hiếm có trọng số ngang lớp phổ biến** — đúng thứ cần cho bài "
            "toán này, và cũng là metric của bảng xếp hạng.",
            "**Cohen's κ** đo mức đồng thuận sau khi trừ phần trùng do may rủi. Đây là "
            "metric xếp hạng thứ hai của challenge.",
            "**Balanced accuracy** và **accuracy** báo kèm để đối chiếu.",
            "**ECE, MCE, Brier, NLL** cho độ tin cậy của xác suất; **AURC** và macro-F1 "
            "tại các mức coverage cho cơ chế từ chối.",
        ],
    ),
    ("h", "3.8.2.", "Khoảng tin cậy bootstrap và so cặp", 3),
    (
        "p",
        "**Mọi con số đều kèm khoảng tin cậy 95%**, tính bằng bootstrap ở mức bệnh nhân, "
        "phân tầng theo lớp, 2.000 lượt. Không bao giờ báo kết quả tốt nhất trong nhiều "
        "seed, và không bao giờ báo một điểm trần.",
    ),
    (
        "p",
        "Khi so hai cấu hình, phép so dùng **bootstrap ghép cặp trên hiệu** — cùng bệnh "
        "nhân, cùng lượt lấy mẫu — chứ không so hai khoảng tin cậy rời nhau. Lý do: hai "
        "cấu hình đánh giá trên đúng cùng một tập ca nên phần lớn phương sai là chung và "
        "sẽ triệt tiêu trong hiệu; so hai khoảng riêng lẻ bỏ mất phần đó và cho một phép "
        "kiểm yếu hơn thực tế.",
    ),
    ("h", "3.8.3.", "Calibration", 3),
    (
        "p",
        "ECE là trung bình có trọng số của chênh lệch giữa độ tự tin và độ chính xác thật "
        "trên các khoảng tự tin; MCE là chênh lệch lớn nhất. Temperature scaling chia "
        "logit cho một hằng số dương duy nhất — nó **không đổi thứ hạng lớp nên không đổi "
        "macro-F1**, tức calibration là thứ nhận được mà không đánh đổi độ chính xác.",
    ),
    (
        "p",
        "Một chi tiết ảnh hưởng tới con số nhiều hơn dự kiến: **nhiệt độ tối ưu theo NLL "
        "khác hẳn nhiệt độ tối ưu theo ECE**, và chọn nhầm sẽ đẩy mô hình sang thiếu tự "
        "tin. Cả hai đều được báo.",
    ),
    ("h", "3.8.4.", "Selective prediction", 3),
    (
        "p",
        "Đường risk–coverage cho biết sai số trên phần ca mà mô hình chịu tự quyết, khi "
        "coverage giảm dần theo thứ tự tin cậy. AURC là diện tích dưới đường đó — càng "
        "thấp càng tốt — và luôn phải đọc kèm hai mốc: xếp ngẫu nhiên và xếp hoàn hảo, vì "
        "không có chúng thì một đường đi xuống lúc nào cũng trông có vẻ tốt.",
    ),
    (
        "p",
        "Hai cách xếp hạng được so: **xác suất cao nhất** của ensemble, và **mức bất đồng "
        "giữa 5 mô hình**. Dự đoán và mô hình giữ nguyên ở cả hai, chỉ đổi thứ tự từ chối — "
        "đó là điều làm cho phép so này sạch.",
    ),
    (
        "callout",
        "⚠️ **Một metric đã phải đổi giữa dự án, và lý do đáng ghi lại.** Mục tiêu ban đầu "
        "là macro-F1 tại coverage 80%. Ở cỡ mẫu một fold (n ≈ 82) nó không tính được có "
        "nghĩa: tại coverage 50%, một lớp hiếm chỉ còn một hai ca, nên F1 của lớp đó do "
        "một bệnh nhân quyết định rồi chiếm 1/7 trọng số macro. Quan sát thực tế: macro-F1 "
        "nhảy loạn trong khi accuracy tăng đều. Metric chính của phần selective vì thế đổi "
        "sang risk–coverage và AURC, và phải tính trên tập gộp out-of-fold chứ không trên "
        "một fold.",
    ),
    ("h", "3.9.", "Hạ tầng và ràng buộc compute", 2),
    (
        "p",
        "Toàn bộ huấn luyện chạy trên Kaggle Notebook với GPU Tesla T4. Ràng buộc "
        "phiên tối đa 12 giờ và quota 30 giờ mỗi tuần định hình mã ngay từ đầu chứ không "
        "phải sửa sau: **checkpoint và resume mỗi epoch**, log ghi thẳng ra CSV không đệm, "
        "mọi đường ghi đi qua biến cấu hình, và AMP bật mặc định.",
    ),
    (
        "p",
        "Cấu hình chính đo thật **0,869 s/batch · 78 s/epoch · 6,50 giờ mỗi fold**, tức đủ "
        "5 fold cần khoảng 32,5 giờ và phải trải qua hai tuần quota. Một khoá cấu hình rẻ "
        "hơn 2–3 lần đã được cân nhắc rồi **bị bác bỏ có chủ ý**: tái lập trung thực recipe "
        "đạt 0,8078 quan trọng hơn tiết kiệm quota, và 32,5 giờ là bài toán kế hoạch chứ "
        "không phải lý do đổi kiến trúc.",
    ),
    (
        "p",
        "Mỗi thí nghiệm đi qua một loạt **cổng kiểm tra chạy trước khi cam kết giờ GPU** — "
        "kiểm trọng số nạp đúng, kiểm hình học cache khớp cấu hình, đo thời gian mỗi epoch "
        "thật, kiểm cán cân dự đoán sau fold đầu. Cổng đo thời gian đã chặn được một cấu "
        "hình ngốn 23,5 giờ ngay lần đầu, mà nguyên nhân chỉ là tăng cường dữ liệu chạy "
        "trên CPU trong khi GPU ngồi chờ.",
    ),
    ("h", "3.10.", "Web app demo", 2),
    (
        "p",
        "Sản phẩm demo là một web app **tự viết full-stack** — FastAPI ở backend, React + "
        "Vite + TypeScript ở frontend — chứ không dùng khung demo dựng sẵn. Giao diện là "
        "một bàn đọc MRI ba cột: dữ liệu đầu vào bên trái, ảnh ở trung tâm, kết quả mô "
        "hình bên phải. Có hai theme sáng và tối, nhưng **vùng ảnh luôn giữ nền đen** kể "
        "cả khi giao diện đang ở theme sáng.",
    ),
    (
        "p",
        "Luồng làm việc một chiều: người dùng thả một tệp ZIP, hệ thống kiểm đủ tám chuỗi "
        "MRI và tám nhãn vùng tổn thương tương ứng, rồi chạy ensemble 5 mô hình ngay trên "
        "máy chủ. Sau khi có kết quả, người dùng xem được đủ tám thì ảnh gốc và bật tắt "
        "nhãn tổn thương. Mức bất định và trạng thái kết quả được diễn đạt bằng ngôn ngữ "
        "thường chứ không bằng thuật ngữ thống kê, và dòng Research Use Only hiện trên mọi "
        "màn hình có kết quả.",
    ),
    (
        "fig",
        "Bàn đọc MRI của web app sau khi xử lý một bộ ảnh đa thì.",
        "w4-webapp-current-mri.png",
        15.5,
    ),
    ("pagebreak",),
    # ================================ 4. RESULT =================================
    ("h", "4.", "Result", 1),
    ("h", "4.1.", "Chuỗi thí nghiệm sàng lọc", 2),
    (
        "p",
        "Lần huấn luyện đầu tiên cho macro-F1 validation **0,2725**. Ba lần thử sửa bằng "
        "siêu tham số đều thất bại, con số đứng yên quanh 0,26–0,27. Thay vì đoán tiếp, dự "
        "án chuyển sang tái lập nguyên khối recipe của baseline chính thức — và bảng đối "
        "chiếu cho thấy sai khác lớn hơn nhiều so với hình dung ban đầu, trong đó weight "
        "decay lệch 5.000 lần.",
    ),
    (
        "p",
        "Từ đó, toàn bộ mức tăng tiếp theo đến từ **cách chuẩn bị dữ liệu, không phải từ "
        "siêu tham số hay kiến trúc**: cắt patch bám sát tổn thương, rồi căn từng thì về "
        "tổn thương của chính nó.",
    ),
    (
        "table",
        "Chuỗi thí nghiệm sàng lọc, đo trên validation fold 1 (82 bệnh nhân, 1 seed).",
        ["Mã", "Thay đổi", "macro-F1 [95% CI]", "κ", "AURC", "ECE thô → sau T"],
        [
            [
                "E0",
                "cửa sổ mm cố định 96×96×48",
                "0,4244 [0,314; 0,530]",
                "0,276",
                "0,5395",
                "0,3218 → 0,1455",
            ],
            [
                "E1",
                "cắt bám sát tổn thương",
                "0,5740 [0,455; 0,678]",
                "0,520",
                "0,2753",
                "0,2935 → 0,2505",
            ],
            ["E3", "hình học 112×112×32", "0,5566 *", "—", "—", "—"],
            [
                "**E4**",
                "**căn từng thì**",
                "**0,7001 [0,599; 0,793]**",
                "**0,646**",
                "**0,2033**",
                "0,2458 → 0,1489",
            ],
        ],
    ),
    (
        "p",
        "* E3 bị **chủ động dừng ở epoch 145/300** nên không dùng làm đối chứng được: "
        "cả ba lần chạy hết đều đạt đỉnh sau epoch 145 (162, 200, 231), nên 0,5566 là cận "
        "dưới chứ không phải trần của cấu hình đó.",
    ),
    (
        "table",
        "So cặp giữa các lần sàng lọc (bootstrap trên hiệu, phân tầng, 2.000 lượt).",
        ["Phép so", "Hiệu macro-F1", "95% CI", "P"],
        [
            ["E4 − E1", "**+0,1261**", "[+0,033; +0,230]", "**0,009**"],
            ["E4 − E0", "+0,2757", "[+0,145; +0,415]", "< 0,001"],
            ["E1 − E0", "+0,1496", "[+0,007; +0,289]", "0,040"],
        ],
    ),
    (
        "callout",
        "⚠️ **Mức tăng +0,126 là chắc chắn; việc quy nó cho phép căn thì chưa.** E4 khác E1 "
        "ở *hai* khoá — hình học và phép căn. Phép so một biến lẽ ra là E4 với E3, nhưng E3 "
        "đã bị dừng sớm nên không gánh được vai trò đó. Báo cáo vì thế quy mức tăng cho "
        "**cả cụm** thay đổi, không tách riêng đóng góp của phép căn.",
    ),
    (
        "p",
        "Ba chỉ báo cơ chế đi kèm đều trúng, và chúng giải thích luôn chứng học thuộc kinh "
        "niên bị ghi nhận suốt E0–E3: hàm mất mát trên validation chạm đáy ở **epoch 9** "
        "với E1 so với **epoch 100** với E4; khoảng cách train/validation cuối +2,55 so với "
        "+1,50; và NLL thô đi từ chỗ *tệ hơn đoán mò* (3,32 so với mức 1,946) xuống 1,72. "
        "Nói cách khác, nguyên nhân học thuộc không nằm ở recipe huấn luyện mà ở đầu vào: "
        "khi tám thì không khớp nhau tới từng voxel thì lớp tích chập đầu tiên không có đặc "
        "trưng liên-thì nào để học, nên nó quay sang ghi nhớ.",
    ),
    ("h", "4.2.", "Cross-validation của cấu hình chính", 2),
    (
        "p",
        "Năm lần huấn luyện, mỗi lần 300 epoch, cùng seed, cấu hình giống hệt nhau trừ chỉ "
        "số fold.",
    ),
    (
        "table",
        "Cross-validation 5 fold của UniFormer-S + Kinetics-400.",
        ["Fold", "n val", "macro-F1", "κ", "Epoch tốt nhất"],
        [
            ["1", "82", "0,8111", "0,7664", "259"],
            ["2", "80", "0,8196", "0,8304", "270"],
            ["3", "78", "0,8293", "0,8238", "103"],
            ["4", "77", "0,7496", "0,7474", "194"],
            ["5", "77", "0,8524", "0,8397", "176"],
            [
                "**Gộp out-of-fold**",
                "**394**",
                "**0,8147** [0,7746; 0,8547]",
                "**0,8010** [0,7600; 0,8418]",
                "—",
            ],
        ],
    ),
    (
        "p",
        "Trung bình năm fold là 0,8124 ± 0,0383. **Con số báo cáo là bản gộp out-of-fold, "
        "không phải trung bình này** — trung bình các fold không có khoảng tin cậy đúng "
        "nghĩa khi mỗi fold là một tập nhỏ khác nhau.",
    ),
    (
        "callout",
        "⚠️ **Thiên lệch do cách chọn checkpoint: +0,0797 [+0,0419; +0,1213].** Checkpoint "
        "được chọn theo macro-F1 trên chính tập validation đang báo, nên 0,8147 **lệch lạc "
        "quan** khoảng 0,08. Con số này được đo và ghi nhận *trước* khi đánh giá trên tập "
        "test, không phải giải thích thêm sau khi thấy kết quả test. Mức thiên lệch tương "
        "ứng của cấu hình đối chứng là +0,079, tức gần bằng nhau.",
    ),
    ("h", "4.3.", "Đánh giá trên tập test khoá kín", 2),
    (
        "p",
        "Tập test 104 ca được giữ kín trong suốt quá trình phát triển. Trước mỗi lần đánh "
        "giá, toàn bộ lựa chọn được ghi thành văn bản và **commit trước khi chạy**: cấu "
        "hình mô hình, bộ dự đoán là ensemble 5 fold, không dùng test-time augmentation, "
        "nhiệt độ fit trên out-of-fold rồi áp mù, danh sách metric, các mức coverage. Tệp "
        "kết quả ghi lại đúng mã commit của bản đăng ký đó, nên quan hệ *protocol có trước "
        "kết quả* là **kiểm được**, không phải một lời hứa.",
    ),
    (
        "callout",
        "⚠️ **Tập test đã được chạm HAI lần, và mọi con số dưới đây phải đọc kèm điều đó.** "
        "Lần một ngày 07/08/2026 với cấu hình DenseNet121-3D; lần hai ngày 14/08/2026 với "
        "cấu hình UniFormer sau khi khoá một bản đăng ký protocol mới. Việc đọc lại tệp xác "
        "suất đã lưu của lần một để so sánh **không** tính là một lần chạm mới, vì không có "
        "suy luận nào chạy.",
    ),
    (
        "table",
        "Kết quả trên tập test 104 ca, hai lần đánh giá. Bộ dự đoán là ensemble 5 fold.",
        ["Lần", "Cấu hình", "macro-F1 [95% CI]", "κ", "Bal. acc", "Accuracy"],
        [
            ["1", "DenseNet121-3D", "0,6162 [0,5246; 0,7032]", "0,5647", "0,6336", "0,6346"],
            [
                "**2**",
                "**UniFormer-S + Kinetics**",
                "**0,7682 [0,6902; 0,8422]**",
                "**0,7333**",
                "**0,7822**",
                "**0,7788**",
            ],
        ],
    ),
    (
        "p",
        "**Mức hụt từ out-of-fold sang test khớp với thiên lệch đã đo trước.** Cấu hình đối "
        "chứng đi từ 0,6851 xuống 0,6162, hụt 0,069, trong khi thiên lệch chọn checkpoint "
        "đo được là +0,079 — hai con số gần trùng khít, nghĩa là phần lạc quan của "
        "out-of-fold đúng bằng phần dự án đã tự chỉ ra, không có nguồn thổi phồng nào khác "
        "lộ ra. Cấu hình chính hụt **ít hơn**: 0,8147 xuống 0,7682, tức −0,047, nên nó "
        "không chỉ điểm cao hơn mà còn khái quát hoá tốt hơn.",
    ),
    (
        "p",
        "**Ensemble lần này có tác dụng thật.** Hiệu so với trung bình 5 mô hình đơn là "
        "**+0,0380 [+0,0007; +0,0771], P = 0,048**, và ensemble vượt cả năm thành viên "
        "(trung bình 0,7302 ± 0,0278, tốt nhất 0,7569). Ở lần chạm một thì ngược lại: hiệu "
        "chỉ +0,0162 với P = 0,43, và mô hình đơn tốt nhất còn cao hơn ensemble. Con số của "
        "mô hình đơn tốt nhất **không được dùng làm kết quả** ở cả hai lần — chọn nó sau "
        "khi nhìn tập test là chọn trên tập test.",
    ),
    ("h", "4.4.", "So sánh có kiểm soát với cấu hình đối chứng", 2),
    (
        "p",
        "Phép so dưới đây hợp lệ vì cấu hình chính được chọn **hoàn toàn trên dữ liệu "
        "out-of-fold**, không dùng một thông tin nào của tập test.",
    ),
    (
        "table",
        "Bootstrap ghép cặp trên cùng bệnh nhân, phân tầng theo lớp, 2.000 lượt.",
        ["Phép so", "Tập đánh giá", "Hiệu macro-F1", "95% CI", "P"],
        [
            [
                "UniFormer − DenseNet121-3D",
                "out-of-fold, 394 ca",
                "**+0,1296**",
                "[+0,0778; +0,1809]",
                "**< 0,001**",
            ],
            [
                "UniFormer − DenseNet121-3D",
                "test, 104 ca",
                "**+0,1520**",
                "[+0,0647; +0,2421]",
                "**0,001**",
            ],
            [
                "Ensemble − trung bình 5 mô hình đơn",
                "test, 104 ca",
                "+0,0380",
                "[+0,0007; +0,0771]",
                "0,048",
            ],
        ],
    ),
    (
        "p",
        "Trên out-of-fold, **cả 5 fold đều dương** (từ +0,082 đến +0,191) và **cả 7 lớp đều "
        "dương** (từ +0,081 đến +0,212). Đây là can thiệp duy nhất trong cả dự án vượt cấu "
        "hình đối chứng có ý nghĩa thống kê, và cũng là lần đầu một hiệu ứng **mạnh lên** "
        "khi tăng cỡ mẫu: đo riêng fold 1 cho +0,111, gộp đủ 394 ca cho +0,130.",
    ),
    ("h", "4.5.", "Vị trí so với văn liệu", 2),
    (
        "table",
        "Các phương pháp đo trên cùng tập test 104 ca chính thức.",
        ["Phương pháp", "macro-F1", "κ"],
        [
            ["Hạng 1 challenge", "0,8322", "0,7801"],
            ["CGHNet (2026)", "0,8180", "0,7820"],
            ["Hạng 2 challenge", "0,8078", "0,7660"],
            ["STM-Former", "0,7930", "0,7520"],
            ["Hạng 3 challenge", "0,7860", "0,7435"],
            ["Hạng 4 challenge", "0,7807", "0,7312"],
            ["**Nghiên cứu này** (ensemble 5 fold)", "**0,7682** [0,6902; 0,8422]", "**0,7333**"],
            ["Hạng 5 challenge", "0,7609", "0,7084"],
            ["DenseNet121-3D (lần chạm 1)", "0,6162", "0,5647"],
            ["Baseline ban tổ chức", "0,6083", "0,5414"],
        ],
    ),
    (
        "p",
        "**Câu phát biểu được phép viết:** kết quả này vượt baseline chính thức 0,6083 một "
        "cách có ý nghĩa thống kê — khoảng tin cậy [0,690; 0,842] không chứa con số đó. "
        "Đây là điều lần chạm một **không** nói được, vì hồi đó 0,6162 có khoảng tin cậy "
        "phủ trùm 0,6083 rất thoải mái.",
    ),
    (
        "callout",
        "⛔ **Câu không được phép viết:** “ngang đội hạng hai”, “ngang "
        "CGHNet”, “tiệm cận SOTA”. Khoảng tin cậy rộng ±0,09 nên không loại "
        "được bất kỳ mốc nào từ 0,709 trở lên. **Định vị đúng:** trên baseline chính thức "
        "và trên nhóm cuối bảng xếp hạng, dưới các phương pháp công bố tốt nhất, và ở "
        "n = 104 thì chưa phân biệt được với cả nhóm 0,71–0,83.",
    ),
    ("h", "4.6.", "Độ tin cậy của xác suất", 2),
    (
        "p",
        "Đây là nửa thứ nhất của đóng góp chính. Nhiệt độ được fit trên 394 ca out-of-fold "
        "rồi áp mù lên tập test, không bao giờ fit trên test.",
    ),
    (
        "table",
        "Calibration của ensemble 5 fold trên tập test (accuracy thật 0,7788).",
        ["Cấu hình xác suất", "ECE", "MCE", "Brier", "NLL", "Tự tin TB (lệch)"],
        [
            [
                "**Chưa hiệu chỉnh** (số chính)",
                "**0,0833**",
                "0,3716",
                "0,3075",
                "0,6804",
                "0,820 (**+0,042**)",
            ],
            [
                "Temperature scaling, T = 1,35",
                "0,0985",
                "**0,2384**",
                "0,3025",
                "0,6656",
                "0,749 (−0,030)",
            ],
        ],
    ),
    (
        "fig",
        "Reliability diagram của ensemble 5 fold trên tập test, xác suất chưa hiệu chỉnh.",
        "fig-reliability.png",
        11.5,
    ),
    ("p", "Ba nhận xét, theo thứ tự quan trọng:"),
    (
        "ul",
        [
            "**ECE 0,0833 đạt được mà không hiệu chỉnh gì**, so với 0,1303 của ensemble ở "
            "lần chạm một và 0,1534 của mô hình đơn tốt nhất *sau khi đã* hiệu chỉnh. Đây "
            "là con số mạnh nhất của phần đóng góp chính.",
            "**Mức tự tin thái quá chỉ còn +0,042**, so với +0,115 ở lần chạm một và +0,186 "
            "trên out-of-fold của cấu hình đối chứng. Trên out-of-fold của cấu hình chính, "
            "ECE là 0,1073 và nhiệt độ cần thiết chỉ 1,53 — so với 3,26 của cấu hình đối "
            "chứng, tức mô hình gần như đã được hiệu chỉnh sẵn.",
            "**Hiệu chỉnh nhiệt độ làm ECE xấu đi trong khi làm MCE tốt lên.** Đây là một "
            "đánh đổi, không phải một cải thiện. Bản chưa hiệu chỉnh đã được chốt làm số "
            "chính **trước** khi chạy, chứ không phải chọn sau khi nhìn số.",
        ],
    ),
    (
        "p",
        "Điểm cuối đáng nói riêng: **gộp năm mô hình hoá ra là bộ hiệu chỉnh tốt hơn "
        "temperature scaling** ở bài toán này. Nhiệt độ học từ phân bố của *mô hình đơn* mà "
        "áp lên *ensemble* vốn đã bớt tự tin thì hiệu chỉnh quá tay — điều này đã được dự "
        "đoán trong bản đăng ký protocol và đã xảy ra đúng như vậy ở cả hai lần chạm.",
    ),
    ("h", "4.7.", "Cơ chế từ chối ca không chắc", 2),
    (
        "p",
        "Nửa thứ hai của đóng góp chính. Dự đoán và mô hình giữ nguyên ở cả hai dòng dưới "
        "đây; chỉ đổi cách xếp hạng để quyết định từ chối ca nào.",
    ),
    (
        "table",
        "Selective prediction trên tập test. Mốc đối chiếu: xếp ngẫu nhiên AURC 0,1615; "
        "xếp hoàn hảo 0,0265.",
        ["Cách xếp hạng", "AURC", "F1@100%", "F1@90%", "F1@80%", "F1@70%", "Cov. @ sai ≤10%"],
        [
            [
                "**Xác suất cao nhất**",
                "**0,0494**",
                "0,7682",
                "0,8022",
                "**0,8421**",
                "0,8911",
                "**76,9%**",
            ],
            ["Bất đồng giữa 5 mô hình", "0,0562", "0,7682", "0,8082", "0,8194", "0,8685", "70,2%"],
        ],
    ),
    (
        "fig",
        "Đường risk–coverage của ensemble 5 fold trên tập test.",
        "fig-risk-coverage.png",
        12.5,
    ),
    (
        "table",
        "Kiểm định ghép cặp: mỗi mức coverage so với trường hợp không từ chối ca nào.",
        ["Mức coverage", "Hiệu macro-F1", "95% CI", "P"],
        [
            ["90%", "+0,0340", "[+0,0015; +0,0688]", "**0,044**"],
            ["**80%**", "**+0,0739**", "[+0,0126; +0,1360]", "**0,027**"],
            ["70%", "+0,1229", "[+0,0087; +0,2043]", "**0,033**"],
        ],
    ),
    (
        "p",
        "**Phát biểu dùng được cho báo cáo và cho giao diện:** từ chối 20% số ca khó nhất "
        "nâng macro-F1 từ 0,768 lên **0,842** (P = 0,027), và ở mức chấp nhận sai số dưới "
        "10% thì hệ thống tự quyết được **76,9%** số ca. Con số tương ứng ở lần chạm một là "
        "29%.",
    ),
    (
        "callout",
        "⚠️ **Một dự đoán chốt trước đã sai, và điều đó được ghi lại đúng như vậy.** Bản "
        "đăng ký protocol viết: *“dự đoán selective cũng sẽ không đạt ý nghĩa thống kê "
        "ở mức 80%; nếu nó đạt thì dự đoán này sai và phải ghi rõ là sai”*. **Nó đạt, "
        "ở cả ba mức.** Căn cứ của dự đoán sai là hành vi trên out-of-fold, nơi không mức "
        "coverage nào đạt ý nghĩa và không lỗi nào có biên quyết định hẹp. Bài học: hành vi "
        "selective trên out-of-fold **không dự báo được** hành vi trên tập test.",
    ),
    (
        "p",
        "Cách xếp hạng đơn giản hơn — chỉ lấy xác suất cao nhất — **thắng** cách dùng mức "
        "bất đồng, ở cả AURC (hiệu +0,0069 nghiêng về xác suất cao nhất, P = 0,053) lẫn "
        "F1@80% (+0,0235, P = 0,38). Đây là kết luận có giá trị thực dụng: với năm mô hình "
        "độc lập thật, softmax của trung bình đã là tín hiệu bất định đủ tốt, **không cần "
        "MC-dropout hay bất kỳ đại lượng bất đồng nào**.",
    ),
    ("h", "4.8.", "Kết quả từng lớp", 2),
    (
        "table",
        "macro-F1 từng lớp, kèm precision và recall trên out-of-fold.",
        ["Lớp", "n val", "F1 out-of-fold", "n test", "F1 test", "P (OOF)", "R (OOF)"],
        [
            ["u máu", "63", "0,912 [0,857; 0,960]", "16", "0,938 [0,839; 1,000]", "0,919", "0,905"],
            ["nang", "42", "0,897 [0,833; 0,955]", "11", "0,957 [0,880; 1,000]", "0,867", "0,929"],
            ["FNH", "36", "0,895 [0,822; 0,959]", "10", "0,909 [0,800; 1,000]", "0,850", "0,944"],
            ["HCC", "125", "0,878 [0,835; 0,916]", "32", "0,787 [0,667; 0,885]", "0,862", "0,896"],
            ["áp-xe", "42", "0,814 [0,725; 0,894]", "12", "0,667 [0,400; 0,870]", "0,795", "0,833"],
            ["ICC", "46", "0,731 [0,630; 0,821]", "12", "0,621 [0,438; 0,800]", "0,723", "0,739"],
            [
                "**di căn**",
                "40",
                "**0,576** [0,433; 0,714]",
                "11",
                "**0,500** [0,222; 0,737]",
                "0,731",
                "**0,475**",
            ],
        ],
    ),
    ("fig", "Ma trận nhầm lẫn bảy lớp trên tập test.", "fig-confusion.png", 11.5),
    (
        "callout",
        "⚠️ Số ca mỗi lớp trên tập test chỉ 10–16, nên **không diễn giải sâu từng con số** "
        "ở cột đó. Khoảng tin cậy của các lớp hiếm rộng tới ±0,25.",
    ),
    (
        "p",
        "**Chẩn đoán nút thắt đảo chiều so với cấu hình đối chứng.** Cấu hình cũ dự đoán "
        "*thừa* hai lớp yếu — ICC 1,26 lần và áp-xe 1,31 lần so với số ca thật — tức vấn đề "
        "nằm ở precision. Cấu hình mới đã cân sáu lớp, riêng di căn lật hẳn sang *thiếu*: "
        "precision tốt (0,731) nhưng recall chỉ 0,475. Nói cách khác, **mô hình giờ quá dè "
        "dặt khi gọi tên di căn: nói ra thì thường đúng, nhưng bỏ sót hơn một nửa.**",
    ),
    (
        "p",
        "Ba hướng nhầm lớn nhất trên out-of-fold là di căn → HCC (8 ca), ICC → HCC (6), "
        "HCC → ICC (4). Trên tập test, ICC là lớp hút nhầm chính: di căn → ICC 3 ca, "
        "HCC → ICC 3 ca. Đáng chú ý là **lớp đa số không còn là nút thắt** — chữa hết 13 "
        "lỗi của HCC chỉ nâng macro-F1 out-of-fold từ 0,8147 lên 0,8405, nhỏ hơn nhiều so "
        "với mức +0,060 tương ứng ở cấu hình cũ.",
    ),
    (
        "p",
        "**Trần số học:** nếu sáu lớp còn lại đều đạt 0,95 mà di căn giữ nguyên 0,576 thì "
        "macro-F1 cũng chỉ tới **0,896**. Lớp di căn một mình chặn mốc 0,9.",
    ),
    ("h", "4.9.", "Các hướng đã thử không hiệu quả", 2),
    (
        "p",
        "Kết quả âm cũng là kết quả, và ở đây chúng chiếm phần lớn số giờ GPU của dự án.",
    ),
    (
        "table",
        "Các can thiệp không đạt ý nghĩa thống kê. Tất cả dùng bootstrap ghép cặp.",
        ["Thí nghiệm", "Tập đánh giá", "Hiệu macro-F1", "95% CI", "P"],
        [
            ["Focal loss (γ = 2)", "2 fold, 162 ca", "−0,029", "[−0,105; +0,048]", "0,47"],
            ["Tăng cường dữ liệu mạnh hơn", "2 fold, 162 ca", "−0,014", "[−0,078; +0,052]", "0,68"],
            ["Bỏ nhiễu cường độ theo thì", "5 fold, 394 ca", "−0,002", "[−0,042; +0,036]", "0,92"],
            [
                "Test-time augmentation (lật)",
                "5 fold, 394 ca",
                "−0,015",
                "[−0,035; +0,004]",
                "0,15",
            ],
            ["Tái lập CGHNet", "5 fold, 394 ca", "−0,019", "[−0,068; +0,031]", "0,46"],
            ["Gộp DenseNet ⊕ CGHNet", "5 fold, 394 ca", "−0,010", "[−0,039; +0,018]", "0,47"],
        ],
    ),
    (
        "p",
        "Gộp cấu hình chính với cấu hình đối chứng cũng **làm tệ đi ở mọi trọng số đã thử**, "
        "từ 0,7349 khi chia đều tới 0,8129 khi cấu hình chính chiếm 90% — đều thấp hơn "
        "0,8147 của nó đứng một mình. Điều này xảy ra dù tỉ lệ trùng lặp lỗi giữa hai mô "
        "hình chỉ 61% (kỳ vọng 30% nếu độc lập) và mức trần oracle lên tới 0,901. **Trùng "
        "lặp lỗi thấp không bảo đảm ensemble ăn**; muốn khai thác 8,6 điểm dư địa đó thì "
        "cần một bộ phối hợp học được, không phải phép trung bình cố định.",
    ),
    (
        "p",
        "Hai kết quả âm đáng ghi thêm vì chúng nói ra điều gì đó về mô hình chứ không chỉ "
        "về can thiệp:",
    ),
    (
        "ul",
        [
            "**Focal loss không cần thiết, và lý do rất cụ thể.** Nó có làm mô hình bớt tự "
            "tin từ đầu (ECE thô 0,154 so với 0,221), nhưng sau khi hiệu chỉnh đúng cách "
            "thì hai bên bằng nhau: 0,1255 và 0,1281. Lợi thế biến mất qua đúng bước mà "
            "pipeline vốn đã làm.",
            "**Mô hình không bất biến với chính phép tăng cường của nó.** Phép lật được áp "
            "độc lập trên cả ba trục với xác suất 0,5, nên phân bố huấn luyện đối xứng hoàn "
            "toàn với phép lật; vậy mà mô hình vẫn mất 0,02–0,06 điểm khi ảnh bị lật. Đây "
            "là bằng chứng độc lập và sạch nhất cho câu chuyện học thuộc, vì nó đo ở một "
            "checkpoint cố định, không dính gì tới chuyện chọn epoch.",
        ],
    ),
    (
        "callout",
        "⚠️ **Một lỗi thật, phát hiện muộn, tồn tại suốt 12 lần huấn luyện đầu.** Trong "
        "khâu tăng cường dữ liệu, phép tịnh tiến lấp phần trống bằng 0 và phép xoay lấp góc "
        "bằng 0. Hệ quả: gần **100% mẫu huấn luyện mang một dải đen ở rìa, trong khi 0% mẫu "
        "validation có**. Đây là lệch phân bố có hệ thống ở mọi bước huấn luyện. Bản sửa "
        "(E12) cho kết quả null, và cấu hình chính tránh lỗi này bằng cách cắt ngẫu nhiên "
        "từ một khối rộng hơn cộng lấp biên bằng voxel biên.",
    ),
    ("h", "4.10.", "Thời gian xử lý một ca", 2),
    (
        "table",
        "Thời gian xử lý, theo thành phần.",
        ["Thành phần", "Thiết bị", "Thời gian"],
        [
            ["Đọc và tiền xử lý tám chuỗi MRI", "CPU", "3,43 s (trung vị) – 4,74 s (p90)"],
            ["Suy luận, 1 mô hình *", "GPU Tesla T4", "81,7 ms"],
            ["Suy luận, ensemble 5 mô hình *", "GPU Tesla T4", "408,5 ms"],
            ["**Tổng end-to-end, web app**", "CPU laptop", "**≈ 18 – 22 s**"],
            ["**Tổng end-to-end, web app**", "GPU Tesla T4", "**≈ 3,8 – 5,2 s**"],
        ],
    ),
    (
        "p",
        "* Hai dòng suy luận đo **theo lô** (batch 4, AMP bật, 104 ca) trong chính lượt "
        "đánh giá test, và **không** bao gồm đọc cùng tiền xử lý NIfTI. Chúng không phải "
        "thời gian đáp ứng của hệ thống thật — hai dòng end-to-end mới là con số đó.",
    ),
    (
        "p",
        "Tiền xử lý chiếm phần lớn thời gian chờ, nên muốn giảm độ trễ thì phải tối ưu khâu "
        "đó chứ không phải khâu mô hình. Việc dùng ensemble 5 mô hình thay vì 1 chỉ tốn "
        "thêm khoảng 0,33 giây, tức **không phải một ràng buộc thực tế**.",
    ),
    ("pagebreak",),
    # ============================== 5. CONCLUSION ===============================
    ("h", "5.", "Conclusion", 1),
    ("h", "5.1.", "Trả lời câu hỏi nghiên cứu", 2),
    (
        "p",
        "**Có.** Một mô hình phân loại bảy loại tổn thương gan trên MRI đa thì, huấn luyện "
        "trên 394 bệnh nhân, đưa ra được xác suất đáng tin và biết từ chối ca nó không chắc.",
    ),
    (
        "p",
        "Cụ thể, trên tập test 104 ca chính thức được giữ kín và đánh giá theo protocol khoá "
        "trước: macro-F1 **0,7682 [0,6902; 0,8422]**, vượt baseline chính thức 0,6083 một "
        "cách có ý nghĩa thống kê. Xác suất đạt **ECE 0,0833 mà không cần hiệu chỉnh gì**, "
        "với mức tự tin thái quá chỉ +0,042. Cơ chế từ chối nâng macro-F1 từ 0,768 lên "
        "**0,842** khi bỏ 20% số ca khó nhất (P = 0,027), và ở mức chấp nhận sai số dưới "
        "10% hệ thống tự quyết được **76,9%** số ca.",
    ),
    ("h", "5.2.", "Đóng góp chính", 2),
    (
        "ul",
        [
            "**Xác suất được hiệu chỉnh.** ECE 0,0833 trên tập test mà không qua bước hiệu "
            "chỉnh nào, tốt hơn mọi con số trước đó của dự án kể cả sau hiệu chỉnh. Kèm "
            "một phát hiện dùng được: **gộp nhiều mô hình là bộ hiệu chỉnh tốt hơn "
            "temperature scaling** ở bài toán này, và hiệu chỉnh thêm một mô hình vốn đã "
            "gần calibrated là lợi bất cập hại.",
            "**Cơ chế từ chối có tác dụng đã kiểm định.** Mức tăng có ý nghĩa thống kê ở cả "
            "ba mức coverage, kèm dòng đối chứng cho thấy cách xếp hạng đơn giản nhất là "
            "đủ — không cần MC-dropout.",
            "**Một mức hiệu năng phân loại vượt baseline chính thức có ý nghĩa thống kê**, "
            "đạt được bằng cách đổi nguồn khởi tạo trọng số chứ không bằng cách tăng dung "
            "lượng mô hình.",
            "**Phần phương pháp luận**, và đây có thể là đóng góp bền nhất: protocol khoá "
            "trước rồi commit, tập test chỉ chạm theo đăng ký, bootstrap ghép cặp thay cho "
            "so hai khoảng rời, thiên lệch chọn checkpoint được đo và báo cáo thay vì im "
            "lặng, và các dự đoán sai được ghi lại đúng như chúng đã sai.",
        ],
    ),
    ("h", "5.3.", "Mức hoàn thành so với kế hoạch", 2),
    (
        "table",
        "Mức hoàn thành từng mục tiêu.",
        ["Mục tiêu", "Mức hoàn thành", "Bằng chứng"],
        [
            [
                "Pipeline tái lập từ MRI thô đến bảng metric",
                "**Đạt**",
                "chạy lại được từ config + seed",
            ],
            [
                "Split khoá mức bệnh nhân, test chống rò rỉ",
                "**Đạt**",
                "khớp official 7/7 lớp; 640 test xanh",
            ],
            [
                "Cross-validation 5 fold có khoảng tin cậy",
                "**Đạt**",
                "0,8147 [0,7746; 0,8547], n = 394",
            ],
            [
                "Backbone pre-trained so với huấn luyện từ đầu",
                "**Đạt**",
                "+0,1296 [+0,0778; +0,1809], P < 0,001",
            ],
            ["Đánh giá trên tập test khoá kín", "**Đạt**", "0,7682 [0,6902; 0,8422]"],
            ["Xử lý lớp hiếm", "**Đạt**", "trọng số hiệu dụng + lấy mẫu căn bậc hai"],
            ["Calibration", "**Đạt**", "ECE 0,0833 chưa hiệu chỉnh"],
            ["Selective prediction", "**Đạt**", "+0,0739 tại coverage 80%, P = 0,027"],
            ["Web app demo tự viết full-stack", "**Đạt**", "FastAPI + React, chạy ensemble thật"],
            ["Slide và báo cáo", "**Đạt**", "4 báo cáo tuần + slide + tài liệu này"],
            ["Khả năng giải thích", "Một phần", "heatmap độ nhạy trên 4 ca minh hoạ"],
            ["External validation và OOD probe", "**Chưa làm**", "cắt khỏi phạm vi vì thời gian"],
            ["Rigid registration đầy đủ", "**Chưa làm**", "phép căn hiện chỉ khử tịnh tiến"],
        ],
    ),
    # ============================== 6. DISCUSSION ===============================
    ("h", "6.", "Discussion", 1),
    ("h", "6.1.", "Diễn giải kết quả", 2),
    (
        "p",
        "Kết quả trung tâm của dự án không phải con số 0,7682 mà là **hình dạng của đường "
        "đi tới nó**. Mức tăng từ 0,2725 lên 0,7001 đến hoàn toàn từ cách chuẩn bị dữ liệu, "
        "không một dòng kiến trúc nào bị đụng tới. Mức tăng tiếp theo lên 0,8147 đến hoàn "
        "toàn từ nguồn khởi tạo trọng số. Trong khi đó, **bảy hướng chỉnh hàm mất mát, "
        "ngưỡng quyết định và tăng cường dữ liệu đều cho kết quả null**.",
    ),
    (
        "p",
        "Lời giải thích nhất quán với toàn bộ bằng chứng: ràng buộc của bài toán này nằm ở "
        "**biểu diễn đặc trưng**, không ở siêu tham số. Một bộ chẩn đoán chạy trên xác suất "
        "đã lưu đã chỉ ra điều đó trước khi tốn thêm giờ GPU — nó cho thấy lỗi của mô hình "
        "cũ *rất tự tin* (chỉ 1/117 lỗi có biên quyết định dưới 0,10), *có cấu trúc* (74% "
        "lỗi trùng nhau giữa hai cấu hình khác augmentation), và với lớp di căn thì *thông "
        "tin đơn giản không có mặt* (trong 20 ca sai, không ca nào có di căn ở hạng hai). "
        "Ba dấu hiệu đó cùng loại trừ mọi can thiệp ở tầng quyết định, và chỉ chừa lại thứ "
        "đổi được biểu diễn — tức nguồn khởi tạo trọng số.",
    ),
    (
        "p",
        "Việc trọng số học từ **video** lại giúp cho ảnh y tế 3D thoạt nghe khiên cưỡng, "
        "nhưng nó có nghĩa: video và MRI 3D chia sẻ cấu trúc *khối ba chiều có tương quan "
        "mạnh giữa các lát kề nhau*. Đó là thứ pre-training trên ảnh 2D không cung cấp "
        "được, và cũng là lý do các lần thử với trọng số 2D trước đó không tạo được lợi ích "
        "rõ ràng. Cần nhấn mạnh rằng đây là **quy kết cho cả cụm** sáu thay đổi trong "
        "recipe được tái lập, không phải phép thử một biến sạch cho riêng pre-training.",
    ),
    ("h", "6.2.", "Nút thắt lớp di căn", 2),
    (
        "p",
        "Lớp di căn là ràng buộc duy nhất còn lại đáng kể, và nó chặn mốc 0,9 một mình. "
        "Điều đáng chú ý là **bản chất của nút thắt đã đổi chiều** giữa hai cấu hình: từ "
        "chỗ dự đoán thừa (vấn đề precision) sang chỗ dự đoán thiếu (vấn đề recall).",
    ),
    (
        "p",
        "Sự đảo chiều này có hệ quả ngược với hướng dẫn rút ra từ cấu hình cũ. Ở cấu hình "
        "cũ, trọng số lớp và hiệu chỉnh prior bị loại vì chúng đẩy *sai chiều*; với cấu hình "
        "mới, riêng cho lớp di căn thì chúng lại đúng chiều. Nhưng chúng vẫn bị chặn bởi "
        "cùng một bằng chứng như trước: **không lỗi nào có biên quyết định đủ hẹp để một "
        "phép dịch ngưỡng lật được**. Can thiệp còn khớp phải tác động lúc huấn luyện, "
        "không phải lúc suy luận.",
    ),
    (
        "p",
        "Đây không phải giới hạn riêng của nghiên cứu này. **Lớp yếu nhất của CGHNet ở mức "
        "0,818 cũng là di căn**, và bài báo của họ quy cho số mẫu ít cùng biểu hiện hình "
        "ảnh không đồng nhất. Điều đó hợp lý về mặt lâm sàng: tổn thương di căn thừa hưởng "
        "đặc điểm hình ảnh từ khối u nguyên phát, mà khối u nguyên phát thì có thể ở bất kỳ "
        "cơ quan nào — nên “di căn” về bản chất không phải một lớp đồng nhất "
        "giống như nang hay u máu.",
    ),
    ("h", "6.3.", "Bài học phương pháp luận", 2),
    (
        "p",
        "Bài học lớn nhất của dự án là về quy trình chứ không về mô hình: **một phép sàng "
        "trên cỡ mẫu nhỏ chỉ đủ để loại một ý tưởng, không đủ để chọn nó.** Dự án bị đánh "
        "lừa đúng kiểu đó bốn lần:",
    ),
    (
        "ul",
        [
            "Một can thiệp về tăng cường dữ liệu cho **+0,038 trên 2 fold**, rồi **−0,002 "
            "trên đủ 5 fold**. Toàn bộ mức tăng đến từ một fold may mắn — fold đó cho 0,7660, "
            "con số cao nhất dự án từng có ở thời điểm ấy.",
            "Một phép gộp hai mô hình cho **+0,065 trên 1 fold**, rồi **−0,010 trên 5 fold**.",
            "Cấu hình chính, đo trên riêng fold 1, cho ba con số phụ trợ đều **không sống "
            "sót** qua 394 ca — trong đó có một con số 1,000 trông rất thuyết phục lúc viết.",
            "Ngược lại, hiệu ứng thật thì **mạnh lên** khi tăng cỡ mẫu: +0,111 ở fold 1 "
            "thành +0,130 khi gộp đủ 394 ca.",
        ],
    ),
    (
        "p",
        "Bài học thứ hai: **đối chiếu mốc ngoài trước khi debug.** Ba phiên làm việc đã đốt "
        "ba lần chạy GPU để đoán nguyên nhân một con số thấp, mà không hề biết mức nào mới "
        "là *đạt* cho bài toán này. Cả ba chẩn đoán đều sai. Sau khi tra bảng kết quả có "
        "kiểm soát trong văn liệu — nơi một ResNet3D trần đạt 0,709 chỉ nhờ hình học đầu "
        "vào — thì hướng đi trở nên rõ ràng ngay. Luật này về sau được ghi thẳng vào tài "
        "liệu ngữ cảnh của dự án.",
    ),
    (
        "p",
        "Bài học thứ ba: **ghi lại dự đoán trước khi chạy, kể cả khi nó sai — nhất là khi "
        "nó sai.** Bản đăng ký protocol chứa một dự đoán rằng selective prediction sẽ không "
        "đạt ý nghĩa thống kê trên tập test. Dự đoán đó sai, và vì nó được viết ra trước "
        "nên cái sai trở thành một kết quả đọc được: hành vi selective trên out-of-fold "
        "không dự báo được hành vi trên tập test. Nếu không ghi trước, cùng dữ liệu ấy chỉ "
        "cho một câu kể lại thành công.",
    ),
    ("h", "6.4.", "Giới hạn", 2),
    (
        "ul",
        [
            "**Research Use Only.** Chưa có bất kỳ kiểm định lâm sàng nào. Mọi con số trong "
            "tài liệu này đo trên dữ liệu hồi cứu của một dataset nghiên cứu.",
            "**Cỡ mẫu tập test là 104 ca**, cho khoảng tin cậy rộng ±0,09. Không thể phân "
            "biệt được với bất kỳ phương pháp công bố nào trong khoảng 0,71–0,83, và ở mức "
            "từng lớp thì n = 10–16 nên các con số đó chỉ mang tính mô tả.",
            "**Tập test đã được chạm hai lần.** Dù mỗi lần đều có protocol khoá trước, số "
            "lần chạm càng nhiều thì tính “hoàn toàn chưa nhìn thấy” càng yếu đi. "
            "Lần chạm thứ ba sẽ cần một bản đăng ký mới.",
            "**Con số out-of-fold mang thiên lệch chọn checkpoint +0,0797**, đã đo và báo, "
            "nên nó không phải một ước lượng không thiên lệch của khả năng khái quát hoá.",
            "**Phép căn các thì chỉ khử tịnh tiến**, không khử xoay và không khử biến dạng; "
            "registration đầy đủ chưa được làm.",
            "**Chưa có external validation và chưa có OOD probe.** Toàn bộ kết quả nằm "
            "trong một dataset duy nhất, từ một nhóm bệnh viện, nên chưa nói được gì về "
            "khả năng chuyển sang máy chụp khác hay quần thể khác.",
            "**Nhãn mặt nạ** được gán bán tự động và mức can thiệp của con người không được "
            "mô tả rõ; dự án dùng nó để xác định cửa sổ cắt, nên một sai lệch hệ thống ở đó "
            "sẽ truyền vào mọi kết quả.",
            "**Chế độ tất định không cho tái lập tới từng bit** vì một số phép toán 3D "
            "không tất định trên CUDA. Seed cố định cho phép lặp lại *lần huấn luyện*, "
            "không phải lặp lại từng chữ số — thêm một lý do để mọi con số đều kèm khoảng "
            "tin cậy.",
            "**Tái lập recipe của đội hạng hai chưa trọn vẹn** (bốn chỗ lệch ở mục 3.5), "
            "nên khoảng cách 0,039 điểm so với họ không quy được cho một nguyên nhân đơn lẻ.",
        ],
    ),
    ("h", "6.5.", "Hướng phát triển", 2),
    (
        "ul",
        [
            "**Trục hợp nhất đa thì.** Mọi thí nghiệm của dự án đều đưa tám thì vào làm tám "
            "kênh. Một bảng so sánh một biến trong văn liệu cho thấy chuyển sang khung "
            "Siamese dùng chung encoder mang lại **+0,022 đến +0,052 trên cả sáu backbone "
            "được thử**. Đây là trục dự án chưa từng chạm tới và là hướng có cơ sở nhất.",
            "**Bộ phối hợp học được** thay cho phép trung bình xác suất, để khai thác 8,6 "
            "điểm dư địa oracle mà trung bình cố định không lấy được điểm nào.",
            "**Can thiệp ở tầng huấn luyện cho lớp di căn.** Phần chẩn đoán ở mục 6.2 chỉ "
            "ra rằng mọi phép chỉnh ngưỡng lúc suy luận đều bị chặn, nên hướng duy nhất "
            "còn khớp là thứ tác động lúc huấn luyện — chẳng hạn sinh thêm biến thiên cho "
            "riêng các lớp hiếm.",
            "**External validation và OOD probe**, để biết kết quả có chuyển được sang máy "
            "chụp và quần thể khác hay không. Đây là việc quan trọng nhất nếu muốn tiến gần "
            "hơn tới sử dụng thật.",
            "**Registration đầy đủ** thay cho phép căn chỉ khử tịnh tiến hiện tại.",
            "**Tối ưu khâu tiền xử lý**, vì nó chiếm phần lớn thời gian đáp ứng của web app "
            "trong khi phần mô hình gần như không đáng kể.",
        ],
    ),
    # =============================== 7. REFERENCE ===============================
    ("h", "7.", "Reference", 1),
]


REFERENCES: list[str] = [
    "LLD-MMRI2023 Challenge — Liver Lesion Diagnosis Challenge on Multi-phase MRI, "
    "MICCAI 2023. Repo chính thức và bảng xếp hạng test-104: "
    "https://github.com/LMMMEng/LLD-MMRI2023",
    "Lou, J. và cs. SDR-Former: A Siamese Dual-Resolution Transformer for Liver Lesion "
    "Classification Using 3D Multi-Phase Imaging. arXiv:2402.17246.",
    "Li và cs. CGHNet: Cross-Guided 2D–3D Hybrid Network with attention mechanism for "
    "focal liver lesion classification. Computerized Medical Imaging and Graphics 132 "
    "(2026) 102780. doi:10.1016/j.compmedimag.2026.102780",
    "NPUBXY — giải pháp hạng hai LLD-MMRI 2023, recipe được tái lập trong nghiên cứu "
    "này. Repo: https://github.com/ZHEGG/miccai2023",
    "Li, K. và cs. UniFormer: Unifying Convolution and Self-attention for Visual "
    "Recognition. Trọng số pre-trained dùng trong nghiên cứu này: "
    "https://huggingface.co/Sense-X/uniformer_video (uniformer_small_k400_16x8.pth)",
    "Kay, W. và cs. The Kinetics Human Action Video Dataset. arXiv:1705.06950.",
    "Cui, Y. và cs. Class-Balanced Loss Based on Effective Number of Samples. CVPR 2019.",
    "Lin, T.-Y. và cs. Focal Loss for Dense Object Detection. ICCV 2017.",
    "Guo, C. và cs. On Calibration of Modern Neural Networks. ICML 2017 — temperature "
    "scaling và ECE.",
    "El-Yaniv, R. và Wiener, Y. On the Foundations of Noise-free Selective "
    "Classification. Journal of Machine Learning Research 11 (2010) — risk–coverage và "
    "selective prediction.",
    "Cardoso, M. J. và cs. MONAI: An open-source framework for deep learning in "
    "healthcare imaging. arXiv:2211.02701.",
    "Paszke, A. và cs. PyTorch: An Imperative Style, High-Performance Deep Learning "
    "Library. NeurIPS 2019.",
    "Tài liệu nội bộ của dự án: docs/MRI_Classification_Spec_Sheet.md (chốt kỹ thuật), "
    "docs/TEST104_PREREGISTRATION.md (protocol khoá trước mỗi lần chạm tập test), "
    "reports/W1–W4_REPORT.md (báo cáo tiến độ từng tuần).",
]


# --- Dựng tài liệu ------------------------------------------------------------------


def configure_styles(document: Any) -> None:
    """Khổ giấy, lề và font. Chỉ dùng font có sẵn trên Windows, phủ đủ dấu tiếng Việt."""
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.0))

    styles = document.styles
    _set_style_font(styles["Normal"], BODY_FONT, 11)
    normal = styles["Normal"].paragraph_format
    normal.line_spacing = 1.15
    normal.space_after = Pt(6)

    for level, size in ((1, 16), (2, 13), (3, 11.5)):
        style = styles[f"Heading {level}"]
        _set_style_font(style, HEAD_FONT, size, bold=True)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True


def build_footer(document: Any) -> None:
    """RUO bên trái, số trang bên phải, ở mọi trang (AGENTS.md §12)."""
    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        width = section.page_width - section.left_margin - section.right_margin
        paragraph.paragraph_format.tab_stops.add_tab_stop(width)

        run = paragraph.add_run(RUO_FOOTER + "\t")
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED
        _set_run_font(run, HEAD_FONT)
        _add_field(paragraph, " PAGE ")
        for tail in paragraph.runs[1:]:
            tail.font.size = Pt(8)
            _set_run_font(tail, HEAD_FONT)


def build_cover(document: Any) -> None:
    """Trang bìa. Ngày để trống để người nộp tự điền."""
    for _ in range(3):
        document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÁO CÁO KẾT THÚC DỰ ÁN")
    run.bold = True
    run.font.size = Pt(22)
    _set_run_font(run, HEAD_FONT)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Phân loại đa lớp tổn thương gan trên MRI 3D đa thì,\n"
        "với xác suất được hiệu chỉnh và cơ chế từ chối ca không chắc"
    )
    run.font.size = Pt(13)
    _set_run_font(run, HEAD_FONT)

    document.add_paragraph()
    document.add_paragraph()

    meta = [
        ("Người thực hiện", "Hoàng Đức Trường"),
        ("Người hướng dẫn", "Nguyễn Hoàng Bảo Lam"),
        ("Khối", "VSF-KD&VHVMEC-DL&AI"),
        ("Dataset", "LLD-MMRI (MICCAI 2023), n = 498 bệnh nhân"),
        ("Ngày báo cáo", ""),
        ("Trạng thái", "Research Use Only (RUO) — chưa kiểm định lâm sàng"),
    ]
    table = document.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, value) in zip(table.rows, meta, strict=True):
        left = row.cells[0].paragraphs[0].add_run(label)
        left.bold = True
        left.font.size = Pt(11)
        _set_run_font(left, HEAD_FONT)
        right = row.cells[1].paragraphs[0].add_run(value)
        right.font.size = Pt(11)
        _set_run_font(right, BODY_FONT)

    document.add_paragraph()
    warning = document.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = warning.add_run(
        "Tài liệu nghiên cứu. Không dùng để chẩn đoán, điều trị hay thay thế ý kiến của bác sĩ."
    )
    run.italic = True
    run.font.size = Pt(9)


def build_toc(document: Any) -> None:
    """Mục lục dạng field — Word tự điền khi mở file hoặc khi bấm F9."""
    heading = document.add_paragraph()
    run = heading.add_run("MỤC LỤC")
    run.bold = True
    run.font.size = Pt(14)
    _set_run_font(run, HEAD_FONT)

    paragraph = document.add_paragraph()
    _add_field(paragraph, r' TOC \o "1-3" \h \z \u ')

    note = document.add_paragraph()
    run = note.add_run(
        "(Bấm chuột phải vào mục lục rồi chọn Update Field, hoặc F9, để Word điền số trang.)"
    )
    run.italic = True
    run.font.color.rgb = MUTED
    run.font.size = Pt(9)


def build_body(builder: Builder) -> None:
    """Duyệt CONTENT và dựng từng khối."""
    for block in CONTENT:
        kind = block[0]
        if kind == "h":
            builder.heading(block[1], block[2], block[3])
        elif kind == "p":
            builder.para(block[1])
        elif kind == "ul":
            builder.bullets(block[1])
        elif kind == "callout":
            builder.callout(block[1])
        elif kind == "table":
            builder.table(block[1], block[2], block[3])
        elif kind == "fig":
            width = block[3] if len(block) > 3 else 12.0
            builder.figure(block[1], ASSETS / block[2], width)
        elif kind == "pagebreak":
            builder.page_break()
        else:  # pragma: no cover - lỗi lập trình, không phải lỗi người dùng
            raise ValueError(f"Loại khối không biết: {kind!r}")


def build_references(document: Any) -> None:
    """Danh mục tham khảo, đánh số thủ công để không phụ thuộc list style của Word."""
    for index, reference in enumerate(REFERENCES, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.9)
        paragraph.paragraph_format.first_line_indent = Cm(-0.9)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"[{index}]\t{reference}")
        run.font.size = Pt(10)
        _set_run_font(run, BODY_FONT)


def build_document() -> tuple[Any, Builder]:
    """Ráp toàn bộ tài liệu."""
    document = Document()
    configure_styles(document)
    build_footer(document)

    build_cover(document)
    document.add_page_break()
    build_toc(document)
    document.add_page_break()

    builder = Builder(document)
    build_body(builder)
    build_references(document)

    return document, builder


# --- CLI -----------------------------------------------------------------------------


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Dựng báo cáo kết thúc dự án ra file .docx.")
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    parser.add_argument(
        "--force",
        action="store_true",
        help="Cho phép ghi đè file đã tồn tại. Không có cờ này thì script từ chối ghi đè.",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    out: Path = args.out

    if out.exists() and not args.force:
        raise SystemExit(
            f"'{out}' đã tồn tại. Script này dựng lại toàn bộ tài liệu, nên chạy đè sẽ xoá "
            "mọi chỉnh sửa làm trực tiếp trong Word.\n"
            "Muốn dựng lại thì chạy với --force; nếu không, đổi --out."
        )

    out.parent.mkdir(parents=True, exist_ok=True)
    document, builder = build_document()
    document.save(out)

    print(f"Đã ghi: {out}")
    print(
        f"  python-docx {docx.__version__} · {builder.table_no} bảng · "
        f"{builder.figure_no} hình · {len(REFERENCES)} tham khảo"
    )
    if builder.missing_images:
        print("  ⚠ Thiếu hình (đã để khung trống):")
        for path in builder.missing_images:
            print(f"      {path}")
        print("    Sinh ba hình biểu đồ bằng: python scripts/make_final_report_figures.py")
    print("  Mở bằng Word rồi bấm F9 trên mục lục để điền số trang.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
